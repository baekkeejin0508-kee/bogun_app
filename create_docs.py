import os
from docx import Document

def create_eduzip_checklist():
    doc = Document()
    doc.add_heading('에듀집 탑재용 필수기준 체크리스트 (개인정보 있는 웹앱)', 0)
    
    doc.add_paragraph('웹앱 이름: 보건앱\n개인정보 보호책임자: 백기진 교사 (서울의료보건고등학교)')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '점검 항목'
    hdr_cells[1].text = '확인 결과 (Y/N)'
    hdr_cells[2].text = '웹앱 적용 내용'
    
    items = [
        ('개인정보처리방침 공개 여부 (앱 하단 링크 등)', 'Y', '앱 하단 푸터에 개인정보처리방침 링크 버튼을 구현하여 상세 팝업 제공함'),
        ('개인정보 수집 및 이용 목적의 명확성', 'Y', '학생 회원 가입, 진도율 확인, 콘텐츠 제공 목적으로 명시함'),
        ('수집하는 개인정보 항목의 최소화', 'Y', '아이디, 비밀번호, 이름, 학년, 반, 번호만 최소 수집함'),
        ('만 14세 미만 아동의 법정대리인 동의 절차 안내', 'Y', '개인정보처리방침 및 가입 화면에 가정통신문을 통한 보호자 동의 절차 명시함'),
        ('개인정보 파기 기한 및 절차 명시', 'Y', '학년도 종료 시(익년 2월 말) 또는 졸업/진급 시 5일 이내 파기 명시함'),
        ('비밀번호 등 중요 정보의 암호화 저장 여부', 'Y', '비밀번호 일방향 암호화(Hash) 및 전 구간 HTTPS 보안 통신 적용함'),
        ('개인정보 보호책임자 지정 및 연락처 표기', 'Y', '개발 교사(백기진)를 책임자로 지정하고 소속 및 내선번호 표기함'),
        ('이용약관 게시 및 동의 절차 마련', 'Y', '앱 초기 진입 전 이용약관 확인 및 동의 절차(버튼 클릭) 구현함')
    ]
    
    for item, yn, desc in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = yn
        row_cells[2].text = desc
        
    doc.save('에듀집 탑재용 체크리스트.docx')

def create_school_checklist():
    doc = Document()
    doc.add_heading('학교용 필수 기준 체크리스트 (보건앱)', 0)
    doc.add_paragraph('본 체크리스트는 학교에서 보건앱을 도입 및 활용하기 전 확인해야 할 필수 요건들을 점검하기 위한 양식입니다.')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '구분'
    hdr_cells[1].text = '점검 항목'
    hdr_cells[2].text = '확인 사항'
    
    items = [
        ('관련 법령 준수', '개인정보 보호법 등 관련 법령에 따른 개인정보처리방침이 마련되어 있는가?', '확인 완료 (팝업 안내 적용)'),
        ('관련 법령 준수', '서비스 이용약관이 명확하게 작성되어 안내되고 있는가?', '확인 완료 (팝업 안내 적용)'),
        ('사용자 보호', '만 14세 미만 학생 이용 시 법정대리인의 동의를 받고 있는가?', '확인 완료 (가정통신문 활용 명시)'),
        ('사용자 보호', '수집된 학생의 데이터가 안전하게 관리되고 암호화 전송되는가?', '확인 완료 (HTTPS, Hash 암호화)'),
        ('데이터 관리', '학기 종료 등 개인정보 보유 목적 달성 시 즉시 파기되는가?', '확인 완료 (익년 2월 말 파기 명시)'),
        ('책임 소재', '개인정보 보호책임자가 명확히 지정되어 있고 연락 채널이 있는가?', '확인 완료 (백기진 교사 지정)')
    ]
    
    for cat, item, chk in items:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = item
        row_cells[2].text = chk
        
    doc.save('학교용 필수 기준 체크리스트.docx')

if __name__ == '__main__':
    create_eduzip_checklist()
    create_school_checklist()
